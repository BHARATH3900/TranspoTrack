from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def add_heading_with_color(doc, text, level, color_rgb=(76, 175, 80)):
    """Add a colored heading"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color_rgb)
        run.font.bold = True
    return heading

def create_project_report():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Title Page
    title = doc.add_heading('Transport Management System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(76, 175, 80)
    
    subtitle = doc.add_paragraph('Mini Project Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Project Details
    details = doc.add_paragraph()
    details.add_run('Submitted by: ').bold = True
    details.add_run('[Your Name]\n')
    details.add_run('College: ').bold = True
    details.add_run('[Your College Name]\n')
    details.add_run('Department: ').bold = True
    details.add_run('[Your Department]\n')
    details.add_run('Date: ').bold = True
    details.add_run(datetime.datetime.now().strftime('%B %Y'))
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents
    add_heading_with_color(doc, 'Table of Contents', 1)
    toc_items = [
        ('1. Abstract', 3),
        ('2. Introduction', 3),
        ('3. Problem Statement', 3),
        ('4. Objectives', 3),
        ('5. System Architecture', 3),
        ('6. Technology Stack', 3),
        ('7. Features', 3),
        ('8. Database Design', 3),
        ('9. Implementation', 3),
        ('10. Testing', 3),
        ('11. Future Enhancements', 3),
        ('12. Conclusion', 3),
        ('13. References', 3)
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph(style='List Number')
        p.add_run(item)
    
    doc.add_page_break()
    
    # Abstract
    add_heading_with_color(doc, '1. Abstract', 1)
    doc.add_paragraph(
        'The Transport Management System is a comprehensive web-based application designed to streamline '
        'the management of transport operations for companies operating milk vans and other commercial vehicles. '
        'This system eliminates the need for manual paper-based record-keeping by providing a digital platform '
        'for tracking daily transport entries, calculating costs automatically, and generating professional PDF reports. '
        'Built using Flask (Python) for the backend and vanilla JavaScript for the frontend, the application '
        'offers a responsive, user-friendly interface that works seamlessly across all devices including mobile phones, '
        'tablets, and desktop computers. The system supports multiple vehicle management, automatic cost calculations, '
        'and flexible report generation with customizable date ranges.'
    )
    
    doc.add_page_break()
    
    # Introduction
    add_heading_with_color(doc, '2. Introduction', 1)
    doc.add_paragraph(
        'In today\'s fast-paced business environment, efficient management of transport operations is crucial '
        'for companies that rely on vehicle fleets for their daily operations. Traditional paper-based systems, '
        'such as the manual ledger shown in the project requirements, are prone to errors, difficult to maintain, '
        'and lack the flexibility needed for modern business operations.'
    )
    doc.add_paragraph(
        'This project addresses these challenges by developing a modern web application that digitizes the entire '
        'transport record-keeping process. The application allows users to manage multiple vehicles, record daily '
        'transport entries with automatic calculations, and generate professional PDF reports for documentation '
        'and billing purposes.'
    )
    doc.add_paragraph(
        'The system is designed to be accessible from any device with a web browser, making it convenient for '
        'users to update records on-the-go or from the office. The SQLite database ensures data persistence, '
        'while the Flask framework provides a robust and scalable backend architecture.'
    )
    
    doc.add_page_break()
    
    # Problem Statement
    add_heading_with_color(doc, '3. Problem Statement', 1)
    doc.add_paragraph(
        'Companies operating transport vehicles, particularly in the dairy and logistics sectors, face several '
        'challenges with traditional paper-based record-keeping:'
    )
    
    problems = [
        'Manual calculation errors leading to incorrect billing and financial discrepancies',
        'Difficulty in tracking multiple vehicles and their individual performance',
        'Time-consuming process of creating monthly or quarterly reports',
        'Risk of data loss due to physical damage or misplacement of paper records',
        'Limited accessibility - records can only be accessed at a specific location',
        'Inability to generate quick summaries and statistics for decision-making',
        'Environmental impact of paper usage',
        'Difficulty in sharing information with stakeholders or accounting departments'
    ]
    
    for problem in problems:
        doc.add_paragraph(problem, style='List Bullet')
    
    doc.add_paragraph(
        '\nThese challenges necessitate a digital solution that can automate calculations, provide easy access '
        'from multiple devices, ensure data security, and generate professional reports with minimal effort.'
    )
    
    doc.add_page_break()
    
    # Objectives
    add_heading_with_color(doc, '4. Objectives', 1)
    doc.add_paragraph('The primary objectives of this project are:')
    
    objectives = [
        'Develop a web-based application accessible from all devices (mobile, tablet, desktop)',
        'Implement multi-vehicle management with individual tracking for each vehicle',
        'Automate cost calculations based on kilometers driven and predefined rates',
        'Provide flexibility for users to add extra charges for each entry',
        'Generate professional PDF reports with customizable date ranges',
        'Include sender and recipient address fields in PDF reports for formal documentation',
        'Ensure data persistence using SQLite database',
        'Create an intuitive and responsive user interface requiring minimal training',
        'Implement CRUD (Create, Read, Update, Delete) operations for both vehicles and entries',
        'Provide real-time statistics and summaries for quick insights',
        'Ensure the application can be easily deployed and run locally',
        'Design the codebase to be easily modifiable and maintainable'
    ]
    
    for objective in objectives:
        doc.add_paragraph(objective, style='List Bullet')
    
    doc.add_page_break()
    
    # System Architecture
    add_heading_with_color(doc, '5. System Architecture', 1)
    
    add_heading_with_color(doc, '5.1 Architecture Overview', 2, (33, 150, 243))
    doc.add_paragraph(
        'The Transport Management System follows a three-tier architecture pattern:'
    )
    
    doc.add_paragraph('Presentation Layer (Frontend)', style='List Bullet')
    doc.add_paragraph('HTML5, CSS3, and Vanilla JavaScript for user interface')
    doc.add_paragraph('Responsive design for cross-device compatibility')
    
    doc.add_paragraph('Application Layer (Backend)', style='List Bullet')
    doc.add_paragraph('Flask web framework handling HTTP requests')
    doc.add_paragraph('RESTful API endpoints for data operations')
    doc.add_paragraph('PDF generation using ReportLab library')
    
    doc.add_paragraph('Data Layer', style='List Bullet')
    doc.add_paragraph('SQLite database for data persistence')
    doc.add_paragraph('SQLAlchemy ORM for database operations')
    
    add_heading_with_color(doc, '5.2 Application Flow', 2, (33, 150, 243))
    doc.add_paragraph(
        '1. User accesses the web application through a browser\n'
        '2. Frontend sends HTTP requests to Flask backend via RESTful APIs\n'
        '3. Flask processes requests and interacts with SQLite database using SQLAlchemy\n'
        '4. Data is returned to frontend as JSON responses\n'
        '5. JavaScript updates the DOM to display data dynamically\n'
        '6. For PDF generation, backend uses ReportLab to create formatted documents\n'
        '7. Generated PDFs are sent back to client for download'
    )
    
    doc.add_page_break()
    
    # Technology Stack
    add_heading_with_color(doc, '6. Technology Stack', 1)
    
    # Create technology table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology'
    
    tech_data = [
        ('Backend Framework', 'Flask 3.0.0'),
        ('Programming Language', 'Python 3.8+'),
        ('Database', 'SQLite'),
        ('ORM', 'SQLAlchemy 2.0.23'),
        ('PDF Generation', 'ReportLab 4.0.7'),
        ('Frontend', 'HTML5, CSS3, JavaScript (ES6+)'),
        ('HTTP Client', 'Fetch API'),
        ('Design Pattern', 'MVC (Model-View-Controller)'),
        ('Architecture', 'RESTful API'),
        ('Development Server', 'Flask Development Server')
    ]
    
    for component, tech in tech_data:
        row_cells = table.add_row().cells
        row_cells[0].text = component
        row_cells[1].text = tech
    
    doc.add_paragraph()
    
    add_heading_with_color(doc, '6.1 Why These Technologies?', 2, (33, 150, 243))
    
    doc.add_paragraph('Flask: ', style='List Bullet').runs[0].bold = True
    doc.add_paragraph('Lightweight, easy to learn, perfect for small to medium applications')
    
    doc.add_paragraph('SQLite: ', style='List Bullet').runs[0].bold = True
    doc.add_paragraph('Zero-configuration, serverless, perfect for local deployment')
    
    doc.add_paragraph('Vanilla JavaScript: ', style='List Bullet').runs[0].bold = True
    doc.add_paragraph('No framework dependencies, fast loading, full control over code')
    
    doc.add_paragraph('ReportLab: ', style='List Bullet').runs[0].bold = True
    doc.add_paragraph('Professional PDF generation with complete formatting control')
    
    doc.add_page_break()
    
    # Features
    add_heading_with_color(doc, '7. Features', 1)
    
    add_heading_with_color(doc, '7.1 Vehicle Management', 2, (33, 150, 243))
    features_vehicle = [
        'Create multiple vehicles with custom names',
        'Set default rate per kilometer for each vehicle',
        'Edit vehicle information at any time',
        'Delete vehicles (with confirmation to prevent accidental deletion)',
        'Visual dashboard showing all vehicles as cards',
        'Active vehicle highlighting for easy identification'
    ]
    for feature in features_vehicle:
        doc.add_paragraph(feature, style='List Bullet')
    
    add_heading_with_color(doc, '7.2 Entry Management', 2, (33, 150, 243))
    features_entry = [
        'Add daily transport entries with date, route, kilometers, and rate',
        'Automatic calculation of amount (kilometers × rate)',
        'Support for extra charges with automatic total calculation',
        'Edit existing entries with live recalculation',
        'Delete entries with confirmation',
        'Tabular view of all entries sorted by date',
        'Default rate pre-filled from vehicle settings'
    ]
    for feature in features_entry:
        doc.add_paragraph(feature, style='List Bullet')
    
    add_heading_with_color(doc, '7.3 Reporting & Analytics', 2, (33, 150, 243))
    features_report = [
        'Real-time statistics dashboard showing:',
        '  - Total number of entries',
        '  - Total kilometers driven',
        '  - Total extra charges',
        '  - Grand total amount',
        'PDF report generation with custom date range',
        'Professional PDF layout with sender and recipient addresses',
        'Detailed table in PDF with all entries',
        'Summary section in PDF with totals and statistics'
    ]
    for feature in features_report:
        doc.add_paragraph(feature, style='List Bullet')
    
    add_heading_with_color(doc, '7.4 User Experience', 2, (33, 150, 243))
    features_ux = [
        'Responsive design - works on all screen sizes',
        'Modal dialogs for data entry and editing',
        'Form validation to prevent invalid data',
        'Success/error messages for user feedback',
        'Confirmation dialogs for destructive actions',
        'Clean and intuitive interface requiring minimal training',
        'Fast loading times with no external dependencies'
    ]
    for feature in features_ux:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # Database Design
    add_heading_with_color(doc, '8. Database Design', 1)
    
    add_heading_with_color(doc, '8.1 Entity Relationship', 2, (33, 150, 243))
    doc.add_paragraph(
        'The database consists of two main entities with a one-to-many relationship:\n\n'
        'Vehicle (1) ←→ (Many) TransportEntry'
    )
    
    add_heading_with_color(doc, '8.2 Vehicle Table', 2, (33, 150, 243))
    
    table_vehicle = doc.add_table(rows=1, cols=4)
    table_vehicle.style = 'Light Grid Accent 1'
    
    hdr = table_vehicle.rows[0].cells
    hdr[0].text = 'Field'
    hdr[1].text = 'Type'
    hdr[2].text = 'Constraints'
    hdr[3].text = 'Description'
    
    vehicle_fields = [
        ('id', 'Integer', 'PRIMARY KEY', 'Unique identifier'),
        ('name', 'String(100)', 'NOT NULL', 'Vehicle name'),
        ('default_rate', 'Float', 'DEFAULT 0.0', 'Default rate per km'),
        ('created_at', 'DateTime', 'DEFAULT NOW', 'Creation timestamp')
    ]
    
    for field, ftype, constraints, desc in vehicle_fields:
        row = table_vehicle.add_row().cells
        row[0].text = field
        row[1].text = ftype
        row[2].text = constraints
        row[3].text = desc
    
    doc.add_paragraph()
    
    add_heading_with_color(doc, '8.3 TransportEntry Table', 2, (33, 150, 243))
    
    table_entry = doc.add_table(rows=1, cols=4)
    table_entry.style = 'Light Grid Accent 1'
    
    hdr = table_entry.rows[0].cells
    hdr[0].text = 'Field'
    hdr[1].text = 'Type'
    hdr[2].text = 'Constraints'
    hdr[3].text = 'Description'
    
    entry_fields = [
        ('id', 'Integer', 'PRIMARY KEY', 'Unique identifier'),
        ('vehicle_id', 'Integer', 'FOREIGN KEY', 'Reference to Vehicle'),
        ('date', 'Date', 'NOT NULL', 'Entry date'),
        ('route_name', 'String(200)', 'NOT NULL', 'Route name'),
        ('km_driven', 'Float', 'NOT NULL', 'Kilometers driven'),
        ('rate', 'Float', 'NOT NULL', 'Rate per kilometer'),
        ('amount', 'Float', 'NOT NULL', 'Calculated amount'),
        ('extra', 'Float', 'DEFAULT 0.0', 'Extra charges'),
        ('total_amount', 'Float', 'NOT NULL', 'Total amount'),
        ('created_at', 'DateTime', 'DEFAULT NOW', 'Creation timestamp')
    ]
    
    for field, ftype, constraints, desc in entry_fields:
        row = table_entry.add_row().cells
        row[0].text = field
        row[1].text = ftype
        row[2].text = constraints
        row[3].text = desc
    
    doc.add_paragraph()
    
    add_heading_with_color(doc, '8.4 Database Relationships', 2, (33, 150, 243))
    doc.add_paragraph(
        'The relationship between Vehicle and TransportEntry is defined with CASCADE delete, '
        'meaning when a vehicle is deleted, all associated transport entries are also deleted. '
        'This ensures data integrity and prevents orphaned records.'
    )
    
    doc.add_page_break()
    
    # Implementation
    add_heading_with_color(doc, '9. Implementation', 1)
    
    add_heading_with_color(doc, '9.1 Backend Implementation', 2, (33, 150, 243))
    doc.add_paragraph(
        'The backend is implemented using Flask, a micro web framework for Python. '
        'The main application file (app.py) contains:'
    )
    
    impl_backend = [
        'Database model definitions using SQLAlchemy ORM',
        'RESTful API endpoints for CRUD operations',
        'JSON serialization for API responses',
        'Error handling and validation',
        'Integration with PDF generator module'
    ]
    for item in impl_backend:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading_with_color(doc, '9.2 Frontend Implementation', 2, (33, 150, 243))
    doc.add_paragraph(
        'The frontend is built with vanilla JavaScript using modern ES6+ features:'
    )
    
    impl_frontend = [
        'Fetch API for asynchronous HTTP requests',
        'DOM manipulation for dynamic content updates',
        'Event listeners for user interactions',
        'Form validation and data formatting',
        'Modal dialogs for user input',
        'Real-time calculation of amounts',
        'Responsive CSS Grid and Flexbox layouts'
    ]
    for item in impl_frontend:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading_with_color(doc, '9.3 PDF Generation', 2, (33, 150, 243))
    doc.add_paragraph(
        'PDF generation is handled by a separate module (pdf_generator.py) using ReportLab:'
    )
    
    impl_pdf = [
        'Custom page layout with sender/recipient addresses',
        'Professional table styling with alternating row colors',
        'Automatic calculation of totals and summaries',
        'Date range display and vehicle information',
        'Proper formatting with currency symbols and decimal places'
    ]
    for item in impl_pdf:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # Testing
    add_heading_with_color(doc, '10. Testing', 1)
    
    add_heading_with_color(doc, '10.1 Test Cases', 2, (33, 150, 243))
    
    doc.add_paragraph('Vehicle Management Tests:', style='List Bullet')
    doc.add_paragraph('  - Create vehicle with valid data')
    doc.add_paragraph('  - Create vehicle with duplicate name')
    doc.add_paragraph('  - Update vehicle information')
    doc.add_paragraph('  - Delete vehicle with existing entries')
    
    doc.add_paragraph('Entry Management Tests:', style='List Bullet')
    doc.add_paragraph('  - Add entry with all fields')
    doc.add_paragraph('  - Add entry with zero extra charges')
    doc.add_paragraph('  - Update entry and verify recalculation')
    doc.add_paragraph('  - Delete entry and verify removal')
    
    doc.add_paragraph('Calculation Tests:', style='List Bullet')
    doc.add_paragraph('  - Verify amount = km × rate')
    doc.add_paragraph('  - Verify total = amount + extra')
    doc.add_paragraph('  - Test with decimal values')
    
    doc.add_paragraph('PDF Generation Tests:', style='List Bullet')
    doc.add_paragraph('  - Generate PDF with single entry')
    doc.add_paragraph('  - Generate PDF with multiple entries')
    doc.add_paragraph('  - Generate PDF with custom date range')
    doc.add_paragraph('  - Verify all data appears correctly in PDF')
    
    add_heading_with_color(doc, '10.2 Browser Compatibility', 2, (33, 150, 243))
    doc.add_paragraph('The application has been tested on:')
    browsers = [
        'Google Chrome (latest)',
        'Mozilla Firefox (latest)',
        'Microsoft Edge (latest)',
        'Safari (latest)',
        'Mobile browsers (Chrome, Safari on iOS/Android)'
    ]
    for browser in browsers:
        doc.add_paragraph(browser, style='List Bullet')
    
    doc.add_page_break()
    
    # Future Enhancements
    add_heading_with_color(doc, '11. Future Enhancements', 1)
    
    enhancements = [
        'User Authentication and Authorization',
        '  - Multi-user support with login system',
        '  - Role-based access control (admin, user, viewer)',
        '  - User-specific vehicle management',
        '',
        'Advanced Reporting',
        '  - Export to Excel/CSV formats',
        '  - Graphical charts and analytics',
        '  - Monthly/Quarterly comparison reports',
        '  - Fuel consumption tracking',
        '',
        'Cloud Integration',
        '  - Cloud database support (PostgreSQL, MySQL)',
        '  - Automatic backups to cloud storage',
        '  - Remote access from anywhere',
        '',
        'Mobile Application',
        '  - Native Android/iOS applications',
        '  - Offline mode with sync capability',
        '  - Push notifications for reminders',
        '',
        'Additional Features',
        '  - Driver management and assignment',
        '  - Maintenance schedule tracking',
        '  - Route optimization suggestions',
        '  - Automated email reports',
        '  - GPS integration for route tracking',
        '  - Invoice generation from entries',
        '  - Multi-currency support',
        '  - Dark mode theme'
    ]
    
    for enhancement in enhancements:
        if enhancement:
            if enhancement.startswith('  '):
                doc.add_paragraph(enhancement.strip(), style='List Bullet')
            else:
                p = doc.add_paragraph(enhancement)
                p.runs[0].bold = True
        else:
            doc.add_paragraph()
    
    doc.add_page_break()
    
    # Conclusion
    add_heading_with_color(doc, '12. Conclusion', 1)
    doc.add_paragraph(
        'The Transport Management System successfully addresses the challenges of manual paper-based '
        'record-keeping by providing a modern, efficient, and user-friendly digital solution. '
        'The application demonstrates the effective use of web technologies to solve real-world business problems.'
    )
    doc.add_paragraph(
        'Key achievements of this project include:'
    )
    
    achievements = [
        'Complete digitization of transport record-keeping process',
        'Automatic cost calculations eliminating manual errors',
        'Multi-device accessibility improving convenience',
        'Professional PDF report generation for documentation',
        'Intuitive user interface requiring minimal training',
        'Scalable architecture supporting future enhancements',
        'Zero external dependencies making deployment simple'
    ]
    
    for achievement in achievements:
        doc.add_paragraph(achievement, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The modular design and clean codebase make it easy to extend the application with additional features. '
        'The project serves as a solid foundation for further development and can be adapted for use in various '
        'industries beyond dairy transport, including logistics, courier services, and fleet management.'
    )
    doc.add_paragraph(
        'This project demonstrates that with the right technology stack and thoughtful design, complex business '
        'processes can be simplified and made more efficient, ultimately saving time and reducing costs for organizations.'
    )
    
    doc.add_page_break()
    
    # References
    add_heading_with_color(doc, '13. References', 1)
    
    references = [
        'Flask Documentation - https://flask.palletsprojects.com/',
        'SQLAlchemy Documentation - https://www.sqlalchemy.org/',
        'ReportLab Documentation - https://www.reportlab.com/docs/',
        'MDN Web Docs (HTML, CSS, JavaScript) - https://developer.mozilla.org/',
        'Python Official Documentation - https://docs.python.org/3/',
        'SQLite Documentation - https://www.sqlite.org/docs.html',
        'REST API Design Best Practices',
        'Responsive Web Design Principles',
        'Database Normalization Concepts'
    ]
    
    for i, ref in enumerate(references, 1):
        doc.add_paragraph(f'{i}. {ref}')
    
    # Save document
    doc.save('/home/claude/transport_management/PROJECT_REPORT.docx')
    print("Project report generated successfully: PROJECT_REPORT.docx")

if __name__ == '__main__':
    create_project_report()
